import PyPDF2
import docx
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

class FileHandler:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc']
    
    def extract_text(self, file):
        """Extract text from uploaded file (PDF or DOCX)"""
        try:
            filename = file.filename.lower()
            
            if filename.endswith('.pdf'):
                return self._extract_from_pdf(file)
            elif filename.endswith('.docx') or filename.endswith('.doc'):
                return self._extract_from_docx(file)
            else:
                raise ValueError(f"Unsupported file format. Please upload PDF or DOCX files.")
                
        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}")
            raise Exception(f"Failed to process file: {str(e)}")
    
    def _extract_from_pdf(self, file):
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise ValueError("Could not extract text from PDF. Please ensure the PDF contains selectable text.")
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    def _extract_from_docx(self, file):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(BytesIO(file.read()))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("Could not extract text from DOCX file. The document appears to be empty.")
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def validate_file(self, file):
        """Validate uploaded file"""
        if not file or not file.filename:
            return "No file selected"
        
        filename = file.filename.lower()
        if not any(filename.endswith(ext) for ext in self.supported_extensions):
            return f"Unsupported file format. Please upload {', '.join(self.supported_extensions)} files"
        
        # Check file size (max 10MB)
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if size > 10 * 1024 * 1024:  # 10MB
            return "File too large. Please upload files smaller than 10MB"
        
        return None